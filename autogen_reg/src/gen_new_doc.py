import os,sys,re
from docx import Document
from docx.shared import Inches,Pt
from docx.oxml.ns import nsdecls,qn
from docx.oxml import parse_xml

def cell_shd(obj_cell):
    shading_elm = parse_xml(r'<w:shd w:fill="A6A6A6" {}/>'.format(nsdecls('w')))
    obj_cell._tc.get_or_add_tcPr().append(shading_elm)
def cell_font(obj_cell,b_bold):
    run = obj_cell.paragraphs[0].runs[0]
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "宋体")
    run.bold = b_bold

def add_reg_table(obj_doc,creg):
    tbl = obj_doc.add_table(rows=4, cols=0, style='Table Grid')
    for x in range(3):
        tbl.add_column(width=Inches(1.2))
    tbl.add_column(width=Inches(1.5))

    #add const cell
    tbl.cell(0,0).text = '寄存器名'
    tbl.cell(0,2).text = 'offset'
    tbl.cell(1,0).text = 'type'
    tbl.cell(1,2).text = 'spec'
    tbl.cell(2,0).text = 'SW'
    tbl.cell(2,2).text = 'HW'
    tbl.cell(3,0).text = '信号名'
    tbl.cell(3,1).text = '位段定义'
    tbl.cell(3,2).text = 'default'
    tbl.cell(3,3).text = 'comment'
    #cell_style
    for y in range(3):
        cell_shd ( tbl.cell(y,0) )
        cell_shd ( tbl.cell(y,2) )
        cell_font( tbl.cell(y,0),b_bold=1 )
        cell_font( tbl.cell(y,2),b_bold=1 )
    for x in range(4):
        cell_shd ( tbl.cell(3,x) )
        cell_font( tbl.cell(3,x),b_bold=1 )

    #add var cell
    tbl.cell(0,1).text = creg.regname
    tbl.cell(0,3).text = hex(creg.addr)
    tbl.cell(1,1).text = creg.attr_type
    tbl.cell(1,3).text = creg.attr_spec
    tbl.cell(2,1).text = creg.attr_sw
    tbl.cell(2,3).text = creg.attr_hw
    #cell_style
    for y in range(3):
        cell_font( tbl.cell(y,1),b_bold=0 )
        cell_font( tbl.cell(y,3),b_bold=0 )
    #signal
    for aone_signal in creg.list_signal:
        new_row = tbl.add_row()
        for x in range(4):
            new_row.cells[x].text = aone_signal[x]
            cell_font( new_row.cells[x],b_bold=0 )

def add_tbl_sect(hd_num_int,hd_num_frac,creg,document):
    str_hd = '{:d}.{:d} {:s}({:s})'.format(hd_num_int,hd_num_frac,creg.regname,hex(creg.addr))
    document.add_heading(str_hd, 2)
    add_reg_table(document,creg)

def gen_new_table( list_reg=[],fn_doc='demo.docx' ):
    document = Document()

    hd_num_int = 1
    hd_num_frac= 1
    str_hd = '{:d}. 寄存器定义'.format(hd_num_int)
    document.add_heading(str_hd, 1)

    for one_reg in list_reg:
        add_tbl_sect(hd_num_int,hd_num_frac,one_reg,document)
        hd_num_frac += 1

    document.save(fn_doc)

class reg2doc():
    def __init__(self):
        pass
    def run(self,list_reg=[],fn_doc='demo.docx'):
        gen_new_table(list_reg,fn_doc)

if __name__ == "__main__":
    gen_new_table(list_reg)